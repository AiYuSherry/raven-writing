"""Text extractors for local material-library documents."""

import csv
import html
import os
import re
from dataclasses import dataclass, field


@dataclass
class ExtractedDocument:
    title: str
    text: str
    pages: list = field(default_factory=list)
    metadata: dict = field(default_factory=dict)


def extract(path, source_type=None):
    """Extract readable text from a supported local file."""
    ext = os.path.splitext(path)[1].lower()
    kind = (source_type or ext.lstrip(".") or "txt").lower()
    if kind in {"txt", "text", "md", "markdown", "paste"} or ext in {".txt", ".md"}:
        return _extract_plain(path)
    if kind == "html" or ext in {".html", ".htm"}:
        return _extract_html(path)
    if kind == "csv" or ext == ".csv":
        return _extract_csv(path)
    if kind in {"xlsx", "xlsm"} or ext in {".xlsx", ".xlsm"}:
        return _extract_xlsx(path)
    if kind in {"docx", "word"} or ext == ".docx":
        return _extract_docx(path)
    if kind == "pdf" or ext == ".pdf":
        return _extract_pdf(path)
    raise ValueError(f"不支持的素材格式: {ext or source_type}")


def extract_text(title, text):
    title = title or "粘贴文本"
    text = text or ""
    return ExtractedDocument(
        title=title,
        text=text,
        pages=[{"page": 0, "text": text, "section_title": ""}],
        metadata={"source_type": "paste"},
    )


def _extract_plain(path):
    with open(path, "r", encoding="utf-8-sig", errors="replace") as f:
        text = f.read()
    return ExtractedDocument(
        title=os.path.basename(path),
        text=text,
        pages=[{"page": 0, "text": text, "section_title": ""}],
        metadata={"source_type": os.path.splitext(path)[1].lstrip(".") or "txt"},
    )


def _extract_html(path):
    with open(path, "r", encoding="utf-8-sig", errors="replace") as f:
        raw = f.read()
    title = os.path.basename(path)
    try:
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(raw, "html.parser")
        if soup.title and soup.title.string:
            title = soup.title.string.strip()
        for tag in soup(["script", "style", "noscript"]):
            tag.decompose()
        text = soup.get_text("\n")
    except ImportError:
        cleaned = re.sub(r"(?is)<(script|style).*?>.*?</\1>", "\n", raw)
        cleaned = re.sub(r"(?s)<[^>]+>", "\n", cleaned)
        text = html.unescape(cleaned)
    text = _compact_text(text)
    return ExtractedDocument(
        title=title,
        text=text,
        pages=[{"page": 0, "text": text, "section_title": ""}],
        metadata={"source_type": "html"},
    )


def _extract_csv(path):
    rows = []
    with open(path, "r", encoding="utf-8-sig", errors="replace", newline="") as f:
        for row in csv.reader(f):
            rows.append(" | ".join(str(cell).strip() for cell in row))
    text = "\n".join(rows)
    return ExtractedDocument(
        title=os.path.basename(path),
        text=text,
        pages=[{"page": 0, "text": text, "section_title": ""}],
        metadata={"source_type": "csv"},
    )


def _extract_xlsx(path):
    try:
        from openpyxl import load_workbook
    except ImportError as e:
        raise RuntimeError("缺少 openpyxl，无法解析 xlsx/xlsm 文件。请安装 openpyxl。") from e

    workbook = load_workbook(path, read_only=True, data_only=True)
    parts = []
    pages = []
    for idx, sheet in enumerate(workbook.worksheets, start=1):
        lines = [f"## 工作表: {sheet.title}"]
        for row in sheet.iter_rows(values_only=True):
            values = ["" if cell is None else str(cell).replace("\n", " ").strip() for cell in row]
            if any(values):
                lines.append(" | ".join(values))
        sheet_text = "\n".join(lines)
        parts.append(sheet_text)
        pages.append({"page": idx, "text": sheet_text, "section_title": sheet.title})
    workbook.close()
    text = "\n\n".join(parts)
    return ExtractedDocument(title=os.path.basename(path), text=text, pages=pages, metadata={"source_type": "xlsx"})


def _extract_docx(path):
    try:
        from docx import Document
    except ImportError as e:
        raise RuntimeError("缺少 python-docx，无法解析 docx 文件。请安装 python-docx。") from e

    doc = Document(path)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    text = "\n\n".join(parts)
    return ExtractedDocument(
        title=os.path.basename(path),
        text=text,
        pages=[{"page": 0, "text": text, "section_title": ""}],
        metadata={"source_type": "docx"},
    )


def _extract_pdf(path):
    try:
        import fitz

        pages = []
        texts = []
        doc = fitz.open(path)
        for idx, page in enumerate(doc, start=1):
            page_text = page.get_text("text").strip()
            texts.append(page_text)
            pages.append({"page": idx, "text": page_text, "section_title": ""})
        doc.close()
        return ExtractedDocument(
            title=os.path.basename(path),
            text="\n\n".join(texts),
            pages=pages,
            metadata={"source_type": "pdf", "extractor": "pymupdf"},
        )
    except ImportError:
        pass

    try:
        from pypdf import PdfReader
    except ImportError as e:
        raise RuntimeError("缺少 PDF 解析依赖，无法解析 pdf。请安装 pymupdf 或 pypdf。") from e

    reader = PdfReader(path)
    pages = []
    texts = []
    for idx, page in enumerate(reader.pages, start=1):
        page_text = (page.extract_text() or "").strip()
        texts.append(page_text)
        pages.append({"page": idx, "text": page_text, "section_title": ""})
    return ExtractedDocument(
        title=os.path.basename(path),
        text="\n\n".join(texts),
        pages=pages,
        metadata={"source_type": "pdf", "extractor": "pypdf"},
    )


def _compact_text(text):
    lines = [line.strip() for line in (text or "").splitlines()]
    lines = [line for line in lines if line]
    return "\n".join(lines)
